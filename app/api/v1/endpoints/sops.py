"""SOP Builder REST endpoint."""

from __future__ import annotations

import io
import json
import logging
from datetime import date
from pathlib import Path
from typing import Any

import httpx
from docx import Document
from docx.shared import Pt
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse

from app.agents.sop_builder import SOPBuilderAgent
from app.core.config import get_settings
from app.core.security import verify_api_key
from app.models.schemas import SOPRequest, SOPResponse

logger = logging.getLogger(__name__)
settings = get_settings()

router = APIRouter(prefix="/sops", tags=["SOP Builder"])

TEMPLATE_PATH = Path(__file__).parent.parent.parent.parent / "static" / "sop_template.docx"


# ── Helpers ────────────────────────────────────────────────────────────────────

def _safe_str(v: Any, fallback: str = "") -> str:
    if v is None:
        return fallback
    return str(v).strip() or fallback


def _add_paragraph(doc: Document, text: str, style: str = "Normal") -> None:
    """Add a paragraph, skipping if text is empty."""
    if text.strip():
        doc.add_paragraph(text.strip(), style=style)


def _table_set(table: Any, row: int, col: int, text: str) -> None:
    """Safely write to a table cell."""
    try:
        cell = table.cell(row, col)
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    except Exception:
        pass


async def _call_claude(prompt: str) -> str:
    """Call Anthropic API via httpx and return the text response."""
    async with httpx.AsyncClient(timeout=120) as client:
        resp = await client.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "x-api-key": settings.anthropic_api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json",
            },
            json={
                "model": settings.anthropic_model,
                "max_tokens": 4096,
                "messages": [{"role": "user", "content": prompt}],
            },
        )
        resp.raise_for_status()
        data = resp.json()
        return data["content"][0]["text"]


def _build_docx(sop: dict, topic: str, company_name: str, created_by: str) -> bytes:
    """Fill the TPC SOP template with generated content and return bytes."""
    doc = Document(str(TEMPLATE_PATH))

    today_str = date.today().strftime("%m/%d/%Y")
    title    = _safe_str(sop.get("title"), topic)
    purpose  = _safe_str(sop.get("purpose"))
    scope    = sop.get("scope", {}) or {}
    terms    = sop.get("terminology", []) or []
    overview = sop.get("procedure_overview", []) or []
    steps    = sop.get("procedure_details", []) or []
    qc_items = sop.get("quality_control", []) or []
    archiving = _safe_str(sop.get("archiving"))
    appendices = sop.get("appendices", []) or []
    notes    = _safe_str(sop.get("notes_and_tips"))
    approved_by = _safe_str(sop.get("approval", {}).get("approved_by"), created_by)

    # ── Paragraph 0: Title ─────────────────────────────────────────────────────
    for para in doc.paragraphs:
        if "[SOP Title]" in para.text:
            for run in para.runs:
                run.text = run.text.replace("[SOP Title]", title)
            break

    # ── Table 0: Document Details ──────────────────────────────────────────────
    if doc.tables:
        t0 = doc.tables[0]
        _table_set(t0, 0, 1, created_by)
        _table_set(t0, 1, 1, today_str)
        _table_set(t0, 2, 1, created_by)
        _table_set(t0, 3, 1, today_str)
        _table_set(t0, 4, 1, "1.0")

    # ── Purpose paragraph ─────────────────────────────────────────────────────
    for para in doc.paragraphs:
        if "[Write a brief explanation" in para.text:
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = purpose
            else:
                para.add_run(purpose)
            break

    # ── Scope paragraphs ──────────────────────────────────────────────────────
    for para in doc.paragraphs:
        if "Who it Applies To:" in para.text:
            for run in para.runs:
                if "[List the roles" in run.text:
                    run.text = run.text.replace(
                        "[List the roles, departments, or teams this SOP is intended for.]",
                        _safe_str(scope.get("who"), "Accounting team")
                    )
        if "Frequency:" in para.text and "[Indicate how" in para.text:
            for run in para.runs:
                if "[Indicate how" in run.text:
                    run.text = run.text.replace(
                        "[Indicate how often the process is performed, e.g., daily, weekly, annually.]",
                        _safe_str(scope.get("frequency"), "As needed")
                    )
        # Tools placeholders
        tools = scope.get("tools", []) or []
        for idx, placeholder in enumerate(["[Tool 1, e.g., QuickBooks Online]",
                                           "[Tool 2, e.g., W-9 Forms]",
                                           "[Tool 3, e.g., 1099 Wizard]"]):
            if placeholder in para.text:
                replacement = tools[idx] if idx < len(tools) else ""
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, replacement)

    # ── Table 1: Key Terminology ───────────────────────────────────────────────
    if len(doc.tables) > 1 and terms:
        t1 = doc.tables[1]
        # Fill existing rows first (rows 1-3), then add more if needed
        for i, term_item in enumerate(terms[:3]):
            row_idx = i + 1
            if row_idx < len(t1.rows):
                _table_set(t1, row_idx, 0, _safe_str(term_item.get("term")))
                _table_set(t1, row_idx, 1, _safe_str(term_item.get("definition")))
        # Add extra rows for additional terms
        for term_item in terms[3:]:
            row = t1.add_row()
            _table_set(t1, len(t1.rows) - 1, 0, _safe_str(term_item.get("term")))
            _table_set(t1, len(t1.rows) - 1, 1, _safe_str(term_item.get("definition")))

    # ── Table 2: Procedure Overview ────────────────────────────────────────────
    if len(doc.tables) > 2 and overview:
        t2 = doc.tables[2]
        for i, ov in enumerate(overview[:3]):
            row_idx = i + 1
            if row_idx < len(t2.rows):
                _table_set(t2, row_idx, 0, _safe_str(ov.get("step")))
                _table_set(t2, row_idx, 1, _safe_str(ov.get("task")))
        for ov in overview[3:]:
            row = t2.add_row()
            _table_set(t2, len(t2.rows) - 1, 0, _safe_str(ov.get("step")))
            _table_set(t2, len(t2.rows) - 1, 1, _safe_str(ov.get("task")))

    # ── Procedure Details steps ────────────────────────────────────────────────
    # Find the "Procedure Details" heading and insert steps after it
    insert_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == "Heading 2" and "Procedure Details" in para.text:
            insert_idx = i
            break

    if insert_idx is not None and steps:
        # Remove template step placeholders
        paras_to_remove = []
        found_detail_heading = False
        for i, para in enumerate(doc.paragraphs):
            if para.style.name == "Heading 2" and "Procedure Details" in para.text:
                found_detail_heading = True
                continue
            if found_detail_heading:
                if para.style.name == "Heading 2" and "Quality Control" in para.text:
                    break
                if para.text.strip() not in ("", "[Repeat for as many steps as necessary.]"):
                    paras_to_remove.append(para)

        for p in paras_to_remove:
            p_elem = p._element
            p_elem.getparent().remove(p_elem)

        # Find insertion point (after "Procedure Details" heading)
        ref_para = None
        for para in doc.paragraphs:
            if para.style.name == "Heading 2" and "Quality Control" in para.text:
                ref_para = para
                break

        if ref_para:
            ref_elem = ref_para._element
            for step in reversed(steps):
                step_title = _safe_str(step.get("title"), "Step")
                description = _safe_str(step.get("description"))
                subtasks = step.get("subtasks", []) or []
                resources = step.get("resources", {}) or {}

                # Add resources
                if resources.get("guide"):
                    p = doc.add_paragraph(f"Link to Guide: {resources['guide']}", style="Normal")
                    ref_elem.addprevious(p._element)
                if resources.get("video"):
                    p = doc.add_paragraph(f"Link to Video: {resources['video']}", style="Normal")
                    ref_elem.addprevious(p._element)
                if resources:
                    p = doc.add_paragraph("Resources:", style="Normal")
                    ref_elem.addprevious(p._element)

                # Add subtasks
                for subtask in reversed(subtasks):
                    if subtask.strip():
                        p = doc.add_paragraph(subtask.strip(), style="Normal")
                        ref_elem.addprevious(p._element)

                # Add description
                if description:
                    p = doc.add_paragraph(f"Description: {description}", style="Normal")
                    ref_elem.addprevious(p._element)

                # Add step heading
                p = doc.add_paragraph(step_title, style="Heading 3")
                ref_elem.addprevious(p._element)

    # ── Table 3: Quality Control / Troubleshooting ─────────────────────────────
    if len(doc.tables) > 3 and qc_items:
        t3 = doc.tables[3]
        for i, qc in enumerate(qc_items[:3]):
            row_idx = i + 1
            if row_idx < len(t3.rows):
                _table_set(t3, row_idx, 0, _safe_str(qc.get("issue")))
                _table_set(t3, row_idx, 1, _safe_str(qc.get("solution")))
        for qc in qc_items[3:]:
            row = t3.add_row()
            _table_set(t3, len(t3.rows) - 1, 0, _safe_str(qc.get("issue")))
            _table_set(t3, len(t3.rows) - 1, 1, _safe_str(qc.get("solution")))

    # ── Archiving paragraph ────────────────────────────────────────────────────
    for para in doc.paragraphs:
        if "[Detail how to save" in para.text:
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = archiving or "Save all related documents in the client's shared folder."
            break

    # ── Table 4: Approval ──────────────────────────────────────────────────────
    if len(doc.tables) > 4:
        t4 = doc.tables[4]
        _table_set(t4, 1, 0, approved_by)
        _table_set(t4, 1, 1, today_str)

    # ── Appendices ─────────────────────────────────────────────────────────────
    for para in doc.paragraphs:
        if "Appendix A:" in para.text and "[Name of Document" in para.text:
            for run in para.runs:
                if "[Name of Document" in run.text:
                    run.text = run.text.replace(
                        "[Name of Document or Form]",
                        appendices[0] if appendices else "N/A"
                    )
        if "Appendix B:" in para.text and "[Name of Additional" in para.text:
            for run in para.runs:
                if "[Name of Additional" in run.text:
                    run.text = run.text.replace(
                        "[Name of Additional Resource]",
                        appendices[1] if len(appendices) > 1 else "N/A"
                    )
        if "Appendix C:" in para.text and "[Optional" in para.text:
            for run in para.runs:
                if "[Optional" in run.text:
                    run.text = run.text.replace(
                        "[Optional additional details or diagrams]",
                        appendices[2] if len(appendices) > 2 else "N/A"
                    )

    # ── Notes and Tips ────────────────────────────────────────────────────────
    for para in doc.paragraphs:
        if "[Provide any additional advice" in para.text:
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = notes or "Follow each step carefully and double-check your work in QuickBooks Online."
            break

    # ── Remove style notes section (internal template instructions) ────────────
    style_notes_start = False
    to_remove = []
    for para in doc.paragraphs:
        if "Style Notes for Google Docs:" in para.text:
            style_notes_start = True
        if style_notes_start:
            to_remove.append(para)
    for p in to_remove:
        try:
            p._element.getparent().remove(p._element)
        except Exception:
            pass

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Endpoints ──────────────────────────────────────────────────────────────────

@router.post("/build", summary="Generate a complete accounting SOP")
async def build_sop(
    request: SOPRequest,
    _: str = Depends(verify_api_key),
) -> dict:
    agent = SOPBuilderAgent()
    try:
        return await agent.build_sop(request)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@router.post("/build/quick", summary="Generate SOP from a free-text description")
async def build_sop_quick(
    description: str,
    _: str = Depends(verify_api_key),
) -> dict:
    agent = SOPBuilderAgent()
    try:
        return await agent.build_sop_from_description(description)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@router.post("/generate", summary="Generate SOP as .docx using TPC template")
async def generate_sop_docx(
    topic: str = Query(..., description="SOP topic or process name"),
    company_name: str = Query("", description="Client or company name (optional)"),
    created_by: str = Query("Angela – The Profit Catalyst", description="Author name"),
    _: str = Depends(verify_api_key),
) -> StreamingResponse:
    """
    Generate a complete SOP document using The Profit Catalyst template.
    Claude writes all content; the result is a downloadable .docx file.
    """
    if not TEMPLATE_PATH.exists():
        raise HTTPException(status_code=500, detail="SOP template not found on server.")

    if not settings.anthropic_api_key:
        raise HTTPException(status_code=500, detail="Anthropic API key not configured.")

    company_ctx = f" for {company_name}" if company_name else ""

    prompt = f"""You are an expert accounting SOP writer for a bookkeeping firm called The Profit Catalyst.

Generate a complete, detailed SOP{company_ctx} for the following process:

TOPIC: {topic}

Return ONLY a valid JSON object (no markdown, no code fences) with EXACTLY this structure:

{{
  "title": "Full SOP title",
  "purpose": "2-3 sentence explanation of why this process exists and what it achieves.",
  "scope": {{
    "who": "Roles that perform this SOP (e.g., Bookkeeper, Accounting Manager)",
    "frequency": "How often this is performed (e.g., Monthly, Weekly, Per client request)",
    "tools": ["QuickBooks Online", "Tool 2", "Tool 3"]
  }},
  "terminology": [
    {{"term": "Term 1", "definition": "Clear definition"}},
    {{"term": "Term 2", "definition": "Clear definition"}},
    {{"term": "Term 3", "definition": "Clear definition"}}
  ],
  "procedure_overview": [
    {{"step": "Step 1: Title", "task": "One-line description of what happens in this step"}},
    {{"step": "Step 2: Title", "task": "One-line description"}},
    {{"step": "Step 3: Title", "task": "One-line description"}}
  ],
  "procedure_details": [
    {{
      "title": "Step 1: Step Title",
      "description": "Detailed explanation of what this step involves and why.",
      "subtasks": [
        "Specific action 1 with exact QBO navigation (e.g., Go to Accounting > Chart of Accounts)",
        "Specific action 2",
        "Specific action 3"
      ],
      "resources": {{"video": "", "guide": ""}}
    }},
    {{
      "title": "Step 2: Step Title",
      "description": "Detailed explanation.",
      "subtasks": ["Action 1", "Action 2", "Action 3"],
      "resources": {{"video": "", "guide": ""}}
    }}
  ],
  "quality_control": [
    {{"issue": "Common problem 1", "solution": "How to fix or prevent it"}},
    {{"issue": "Common problem 2", "solution": "How to fix or prevent it"}},
    {{"issue": "Common problem 3", "solution": "How to fix or prevent it"}}
  ],
  "archiving": "Where and how to save records. Include retention guidelines.",
  "approval": {{
    "approved_by": "Angela – The Profit Catalyst"
  }},
  "appendices": [
    "Appendix A description",
    "Appendix B description"
  ],
  "notes_and_tips": "Practical tips and best practices for performing this SOP successfully."
}}

Requirements:
- Write at least 4-6 detailed steps in procedure_details
- Each step must have at least 3 specific subtasks with exact actions
- Include real QuickBooks Online navigation paths where applicable
- Be specific and actionable — write for a bookkeeper with 1 year of experience
- Use professional accounting language
- The content must match the topic exactly: {topic}"""

    try:
        raw = await _call_claude(prompt)
    except httpx.HTTPStatusError as exc:
        logger.error("Anthropic API error: %s", exc)
        raise HTTPException(status_code=502, detail=f"AI service error: {exc.response.status_code}") from exc
    except Exception as exc:
        logger.error("SOP generation failed: %s", exc)
        raise HTTPException(status_code=500, detail=f"SOP generation failed: {exc}") from exc

    # Parse JSON — strip any accidental markdown fences
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.split("```", 2)[-1]
        if cleaned.startswith("json"):
            cleaned = cleaned[4:]
        cleaned = cleaned.rsplit("```", 1)[0].strip()

    try:
        sop_data = json.loads(cleaned)
    except json.JSONDecodeError as exc:
        logger.error("JSON parse error. Raw: %s", raw[:500])
        raise HTTPException(status_code=500, detail="AI returned invalid JSON. Please try again.") from exc

    try:
        docx_bytes = _build_docx(sop_data, topic, company_name, created_by)
    except Exception as exc:
        logger.error("DOCX build error: %s", exc)
        raise HTTPException(status_code=500, detail=f"Document generation error: {exc}") from exc

    safe_title = topic[:50].replace(" ", "_").replace("/", "-")
    filename = f"SOP_{safe_title}.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
